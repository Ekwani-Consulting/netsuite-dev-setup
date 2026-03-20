"""Generate Ekwani Consulting Developer Guide as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

TEAL = RGBColor(0x27, 0x85, 0x75)
TEAL_TABLE_HDR = "7CD9C9"
LOGO_PATH = "/Users/ej/.claude-ekw/skills/create-ekw-sol-doc/ekwani_logo.png"
OUTPUT_PATH = "/Users/ej/AI/ekwani-git-process/NetSuite_Developer_Guide.docx"

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

h1 = doc.styles['Heading 1']
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = TEAL
h1.font.name = 'Calibri'
h1.paragraph_format.space_before = Pt(12)
h1.paragraph_format.space_after = Pt(0)

h2 = doc.styles['Heading 2']
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = TEAL
h2.font.name = 'Calibri'

h3 = doc.styles['Heading 3']
h3.font.size = Pt(11)
h3.font.bold = True
h3.font.color.rgb = TEAL
h3.font.name = 'Calibri'


# --- Helpers ---

def make_table(headers, rows, col_widths=None, font_size=Pt(10)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        trPr.append(cant_split)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TEAL_TABLE_HDR}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = font_size
        run.font.name = 'Calibri'

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = font_size
            run.font.name = 'Calibri'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def add_body_bold(bold_text, normal_text):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run = p.add_run(normal_text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def add_code_block(code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(item, tuple):
            run = p.add_run(item[0] + ": ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run = p.add_run(item[1])
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        if isinstance(item, tuple):
            run = p.add_run(item[0] + " ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run = p.add_run(item[1])
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'


def set_header_footer():
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        run = hp.add_run("Ekwani Consulting, LLC")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        run = fp.add_run("\u00a92022 Ekwani Consulting - For Internal Use Only")
        run.font.name = 'Calibri'
        run.font.size = Pt(9)


# --- Cover Page ---

for _ in range(3):
    doc.add_paragraph("")

if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(3.85))

for _ in range(2):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NetSuite Developer Guide")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = TEAL
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SuiteCloud Development Process for Ekwani Consulting")
run.font.size = Pt(14)
run.font.name = 'Calibri'

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("March 20, 2026")
run.font.size = Pt(11)
run.font.name = 'Calibri'

doc.add_page_break()

# --- Document History ---

doc.add_heading("Document History", level=1)
make_table(
    ["Version", "Changes", "Updated By", "Date"],
    [
        ["1.0", "Initial version", "Ekwani Consulting", "March 20, 2026"],
    ],
    col_widths=[0.8, 3.0, 1.8, 1.2],
)
doc.add_page_break()

# --- Table of Contents ---

doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Running the Setup Tool",
    "2. Importing from NetSuite",
    "3. Pushing to NetSuite",
    "4. Git Workflow - Terminal",
    "5. Git Workflow - GitHub Desktop",
    "6. Quick Reference",
    "7. Joining an Existing Project",
    "8. Rollback Guidance",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

doc.add_page_break()

# --- Set header/footer after TOC ---
set_header_footer()

# =============================================
# SECTION 1: Running the Setup Tool
# =============================================
doc.add_heading("1. Running the Setup Tool", level=1)

add_body("The setup tool automates the entire first-time setup process. Follow these steps carefully.")

doc.add_heading("Step 1: Open Your Terminal", level=2)

add_body_bold("On macOS:", "")
add_bullets([
    "Open Finder > Applications > Utilities > Terminal",
    "Or press Cmd + Space, type \"Terminal\", and press Enter",
])

add_body_bold("On Windows:", "")
add_bullets([
    "Press the Windows key, type \"PowerShell\", and press Enter",
    "Or press Windows key + R, type cmd, and press Enter",
])

doc.add_heading("Step 2: Navigate to Your Projects Folder", level=2)

add_body("Type the following and press Enter:")
add_code_block("cd ~/Projects")

add_body("If this folder doesn't exist yet, create it first:")
add_code_block("mkdir ~/Projects\ncd ~/Projects")

doc.add_heading("Step 3: Make Sure You Have GitHub Access", level=2)

add_body("Since the setup tool lives in a private repository, your computer needs to be authenticated with GitHub. Pick one of these options:")

add_body_bold("Option A: GitHub Desktop (recommended for beginners)", "")
add_numbered([
    "Download and install GitHub Desktop from https://desktop.github.com",
    "Sign in with your GitHub account (the one that has access to the Ekwani Consulting organization)",
    "That's it - signing into GitHub Desktop automatically sets up your Git credentials",
])

add_body_bold("Option B: HTTPS + Git Credential Manager", "")
add_numbered([
    "Install Git from https://git-scm.com (on Windows, Git Credential Manager is bundled; on macOS, install via brew install git or Xcode Command Line Tools)",
    "Clone any private Ekwani Consulting repo to trigger the login flow",
    "A browser window will open - sign in with your GitHub account",
    "Your credentials are now cached automatically",
])

add_body_bold("Option C: SSH key (for experienced developers)", "")
add_numbered([
    'Generate an SSH key: ssh-keygen -t ed25519 -C "your.email@company.com"',
    "Add the public key to your GitHub account: Settings > SSH and GPG keys > New SSH key",
    "Paste the contents of ~/.ssh/id_ed25519.pub",
    "Test it works: ssh -T git@github.com",
])

doc.add_heading("Step 4: Run the Setup Tool", level=2)

add_body("Type the following command and press Enter:")
add_code_block("npx github:Ekwani-Consulting/netsuite-dev-setup")

add_body("What to expect:")
add_numbered([
    ("Welcome screen -", "You'll see a banner that says \"Ekwani Consulting - NetSuite Setup\""),
    ("Prerequisite checks -", "The tool checks if Node.js, npm, Git, and SuiteCloud CLI are installed. If SuiteCloud CLI is missing, it will ask to install it - type Y and press Enter."),
    ("Project name -", "Type a name for your project (e.g., my-netsuite-project) and press Enter"),
    ("Directory -", "Press Enter to use the current folder, or type a different path"),
    ("NetSuite login -", "The tool will ask to set up your NetSuite connection. Your web browser will open to the NetSuite login page. Log in, select your role (must have SuiteCloud Developer permissions), and click Allow. Return to your terminal."),
    ("Git setup -", "The tool creates a Git repository and asks for your remote URL. Paste the GitHub repository URL your team lead provided."),
])

doc.add_heading("Troubleshooting", level=2)

make_table(
    ["Problem", "Solution"],
    [
        ["\"npx: command not found\"", "Node.js is not installed. Download from https://nodejs.org"],
        ["\"Permission denied\" or \"Repository not found\"", "Your GitHub account doesn't have access. Ask your team lead to add you to the Ekwani Consulting organization"],
        ["Browser didn't open during NetSuite login", "Copy the URL from the terminal and paste it into your browser manually"],
        ["\"Role doesn't have permissions\"", "Ask your NetSuite admin to enable SuiteCloud Developer permissions on your role"],
    ],
    col_widths=[2.5, 4.0],
    font_size=Pt(9.5),
)

# =============================================
# SECTION 2: Importing from NetSuite
# =============================================
doc.add_page_break()
doc.add_heading("2. Importing from NetSuite", level=1)

add_body("After setup, pull existing scripts and customizations from NetSuite into your local project. Open the terminal in VS Code (Ctrl + ` or Terminal > New Terminal) and make sure you're inside your project folder.")

doc.add_heading("See What Files Exist in NetSuite", level=2)
add_code_block('suitecloud file:list --folder "/SuiteScripts"')
add_body("This lists all script files in your NetSuite account's SuiteScripts folder.")

doc.add_heading("Import a Specific Script File", level=2)
add_code_block('suitecloud file:import --paths "/SuiteScripts/my_script.js"')
add_body("Replace the path with the actual file path from the list above. The file downloads into your project's src/FileCabinet/SuiteScripts/ folder.")

add_body("To import multiple files at once:")
add_code_block('suitecloud file:import --paths "/SuiteScripts/script1.js" "/SuiteScripts/script2.js"')

doc.add_heading("See What Custom Objects Exist", level=2)
add_code_block("suitecloud object:list --type ALL")
add_body("This lists custom records, fields, workflows, and other customizations in your account.")

doc.add_heading("Import All Custom Objects", level=2)
add_code_block('suitecloud object:import --destinationfolder "/Objects" --type ALL --scriptid ALL')
add_body("This downloads all custom objects into your project's src/Objects/ folder as XML files.")

doc.add_heading("Save Your Imports to Git", level=2)
add_body("After importing, always commit your changes:")
add_code_block('git add .\ngit commit -m "Import existing scripts from NetSuite"\ngit push')
add_body("This saves everything to your Git repository so your teammates can see it too.")

# =============================================
# SECTION 3: Pushing to NetSuite
# =============================================
doc.add_page_break()
doc.add_heading("3. Pushing to NetSuite", level=1)

add_body("When you've made changes locally and want to send them to the NetSuite sandbox, you have three options.")

doc.add_heading("Upload a Single Script File", level=2)
add_body("Use this for quick changes to one script:")
add_code_block('suitecloud file:upload --paths "/FileCabinet/SuiteScripts/my_script.js"')
add_body_bold("Important: ", "Notice the path starts with /FileCabinet/ - this is different from importing. Upload paths are relative to your project's src/ folder. Import paths are relative to NetSuite's File Cabinet.")

doc.add_heading("Push a Single Custom Object", level=2)
add_body("Use this when you've changed one custom record, field, or workflow:")
add_code_block("suitecloud object:update --scriptid customrecord_myrecord")
add_body("Replace customrecord_myrecord with the actual script ID of the object you changed.")

doc.add_heading("Deploy the Entire Project", level=2)
add_body("Use this for bigger changes that involve multiple files and objects:")
add_code_block("suitecloud project:deploy")
add_body("This deploys everything listed in your deploy.xml file to the sandbox.")
add_body_bold("Warning: ", "Since everyone shares the same sandbox, let your team know before running project:deploy. A quick message in your team chat prevents surprises.")

# =============================================
# SECTION 4: Git Workflow - Terminal
# =============================================
doc.add_page_break()
doc.add_heading("4. Git Workflow - Terminal", level=1)

add_body("Git keeps track of all your changes and lets you collaborate with teammates without overwriting each other's work. Here's the daily workflow.")

doc.add_heading("1. Create a Branch", level=2)
add_body("Before making changes, create a new branch. Think of it as your own workspace:")
add_code_block("git checkout -b add-customer-validation")
add_body("Replace add-customer-validation with a short description of what you're working on. Use dashes instead of spaces.")

doc.add_heading("2. Make Your Changes", level=2)
add_body("Edit your files in VS Code. When you're done, check what changed:")
add_code_block("git status")
add_body("This shows which files you modified, added, or deleted.")

doc.add_heading("3. Stage Your Changes", level=2)
add_body("Tell Git which files to include in your next save point:")
add_code_block("git add src/FileCabinet/SuiteScripts/my_script.js")
add_body("To stage all changed files at once:")
add_code_block("git add .")

doc.add_heading("4. Commit Your Changes", level=2)
add_body("Save your staged changes with a descriptive message:")
add_code_block('git commit -m "Add customer validation to sales order script"')
add_body('Write your message as if completing the sentence: "This commit will..." Keep it clear and specific.')

doc.add_heading("5. Switch Back to Main", level=2)
add_body("When your work is done and tested, switch to the main branch:")
add_code_block("git checkout main")

doc.add_heading("6. Pull Latest Changes", level=2)
add_body("Get any changes your teammates have pushed while you were working:")
add_code_block("git pull")

doc.add_heading("7. Merge Your Branch", level=2)
add_body("Bring your changes into main:")
add_code_block("git merge add-customer-validation")

doc.add_heading("8. Push to the Remote Repository", level=2)
add_body("Send your changes to GitHub so everyone can see them:")
add_code_block("git push")

doc.add_heading("What to Do If You See a Merge Conflict", level=2)
add_body("Sometimes Git can't automatically combine your changes with a teammate's changes. You'll see a message like:")
add_code_block("CONFLICT (content): Merge conflict in src/FileCabinet/SuiteScripts/my_script.js\nAutomatic merge failed; fix conflicts and then commit the result.")
add_body_bold("Don't panic. Don't force-push. ", "Ask your team lead or a teammate for help. Merge conflicts are normal and easy to resolve with guidance.")

# =============================================
# SECTION 5: Git Workflow - GitHub Desktop
# =============================================
doc.add_page_break()
doc.add_heading("5. Git Workflow - GitHub Desktop", level=1)

add_body("If you prefer a visual interface, GitHub Desktop does the same things as the terminal commands above.")

doc.add_heading("Download GitHub Desktop", level=2)
add_body("If you haven't already, download it from https://desktop.github.com and sign in with your GitHub account. Signing in also sets up Git credentials for your terminal automatically.")

doc.add_heading("1. Open Your Project", level=2)
add_body("File > Add Local Repository and select your project folder. Or if you've already cloned it, it should appear in the left sidebar.")

doc.add_heading("2. Create a Branch", level=2)
add_body("Click the Current Branch dropdown at the top of the window. Click New Branch. Type a name (e.g., add-customer-validation) and click Create Branch.")

doc.add_heading("3. Make Your Changes", level=2)
add_body("Edit files in VS Code as normal. When you switch back to GitHub Desktop, it will show your changes in the left sidebar. You can click on any file to see exactly what changed.")

doc.add_heading("4. Commit Your Changes", level=2)
add_body("At the bottom-left of the window:")
add_numbered([
    'Type a short summary (e.g., "Add customer validation to sales order script")',
    "Click the Commit to [your-branch-name] button",
])

doc.add_heading("5. Switch to Main", level=2)
add_body("Click the Current Branch dropdown at the top and select main.")

doc.add_heading("6. Pull Latest Changes", level=2)
add_body("Click the Fetch origin button at the top. If there are new changes, it will change to Pull origin - click it to download them.")

doc.add_heading("7. Merge Your Branch", level=2)
add_body("Go to Branch menu > Merge into current branch. Select your feature branch from the list and click Merge.")

doc.add_heading("8. Push", level=2)
add_body("Click the Push origin button at the top to send everything to GitHub.")

doc.add_heading("Merge Conflicts in GitHub Desktop", level=2)
add_body_bold("Don't panic. ", "If there's a conflict, GitHub Desktop will show a warning and list the conflicted files. Ask your team lead for help - they can walk you through resolving it in VS Code.")

# =============================================
# SECTION 6: Quick Reference
# =============================================
doc.add_page_break()
doc.add_heading("6. Quick Reference", level=1)

make_table(
    ["What you want to do", "Command"],
    [
        ["List files in NetSuite", 'suitecloud file:list --folder "/SuiteScripts"'],
        ["Import a file", 'suitecloud file:import --paths "/SuiteScripts/file.js"'],
        ["Import all objects", 'suitecloud object:import --destinationfolder "/Objects" --type ALL --scriptid ALL'],
        ["Upload a file", 'suitecloud file:upload --paths "/FileCabinet/SuiteScripts/file.js"'],
        ["Push one object", "suitecloud object:update --scriptid customrecord_name"],
        ["Deploy everything", "suitecloud project:deploy"],
        ["Create a branch", "git checkout -b branch-name"],
        ["Check what changed", "git status"],
        ["Stage all changes", "git add ."],
        ["Stage one file", "git add path/to/file"],
        ["Commit", 'git commit -m "description of changes"'],
        ["Switch to main", "git checkout main"],
        ["Pull latest", "git pull"],
        ["Merge branch", "git merge branch-name"],
        ["Push to GitHub", "git push"],
    ],
    col_widths=[2.0, 4.5],
    font_size=Pt(9.5),
)

# =============================================
# SECTION 7: Joining an Existing Project
# =============================================
doc.add_page_break()
doc.add_heading("7. Joining an Existing Project", level=1)

add_body("If the project already exists and you're joining the team, you don't need the setup tool. Follow these steps instead.")

doc.add_heading("1. Install Prerequisites", level=2)
add_body("Make sure you have these installed:")
add_bullets([
    ("Node.js", "download from https://nodejs.org (LTS version)"),
    ("Git", "download from https://git-scm.com"),
    ("SuiteCloud CLI", "open your terminal and run: npm install -g --acceptSuiteCloudSDKLicense @oracle/suitecloud-cli"),
    ("VS Code", "download from https://code.visualstudio.com"),
    ("SuiteCloud Extension for VS Code", "open VS Code, click the Extensions icon (left sidebar), search \"SuiteCloud\", install the one by Oracle"),
    ("GitHub Desktop (optional)", "download from https://desktop.github.com"),
])

doc.add_heading("2. Clone the Repository", level=2)
add_code_block("git clone https://github.com/ekwani/my-project.git")
add_body("Replace the URL with the actual repository URL your team lead provides.")

doc.add_heading("3. Open the Project", level=2)
add_code_block("cd my-project")
add_body("Then open it in VS Code: File > Open Folder and select the project folder.")

doc.add_heading("4. Install Dependencies", level=2)
add_body("If the project has a package.json file:")
add_code_block("npm install")

doc.add_heading("5. Connect to NetSuite", level=2)
add_body("Each developer needs to set up their own connection:")
add_code_block("suitecloud account:setup")
add_body("This opens your browser for NetSuite login. Follow the prompts to authenticate. This creates a suitecloud.config.js file on your machine only (it's not shared via Git).")

doc.add_heading("6. Verify Your Connection", level=2)
add_code_block('suitecloud file:list --folder "/SuiteScripts"')
add_body("If you see a list of files, you're connected and ready to work.")

# =============================================
# SECTION 8: Rollback Guidance
# =============================================
doc.add_page_break()
doc.add_heading("8. Rollback Guidance", level=1)

doc.add_heading("Something Went Wrong With a Deploy?", level=2)
add_body("If a project:deploy caused issues in the sandbox:")

add_numbered([
    ("Don't panic.", "The previous version of your code is safe in Git."),
    ("Tell your team", "so nobody else deploys on top of the problem."),
    ("Switch to the last known good state:", ""),
])
add_code_block("git checkout main\ngit pull")

add_numbered([
    ("Re-deploy the good version:", ""),
])
add_code_block("suitecloud project:deploy")

add_body("If you're not sure what the \"good version\" is, ask your team lead before deploying anything.")
add_body_bold("Never force-push ", "(git push --force) - this can erase your teammates' work. If Git won't let you push, ask for help.")

# --- Save ---
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
